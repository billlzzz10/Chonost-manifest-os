"""
🎯 RAG System - Retrieval-Augmented Generation.

A comprehensive RAG system for the Synapse Backend Monolith.

Features:
- Vector Database Integration (ChromaDB, Pinecone, Weaviate)
- Multiple Embedding Models (OpenAI, Ollama, Local)
- Document Processing & Chunking
- Semantic Search & Retrieval
- Context-Aware Generation
- Caching & Performance Optimization
"""

import asyncio
import json
import logging
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass, asdict
from pathlib import Path
import numpy as np
import sqlite3
import redis
import requests
from concurrent.futures import ThreadPoolExecutor
import threading

# Vector Database Imports
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False

try:
    import pinecone
    PINECONE_AVAILABLE = True
except ImportError:
    PINECONE_AVAILABLE = False

# AI/ML Imports
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import time
import mimetypes
import magic
import re
import os

# Document Processing Imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

# Code Analysis Imports
try:
    import ast
    import tokenize
    AST_AVAILABLE = True
except ImportError:
    AST_AVAILABLE = False


# Import the unified client from its new shared package location
from packages.ai_clients.unified_ai_client import get_client

@dataclass
class PerformanceMetrics:
    """
    Stores performance metrics.
    """
    start_time: float
    end_time: float
    files_processed: int
    documents_extracted: int
    embeddings_generated: int
    cache_hits: int
    cache_misses: int
    errors: List[str]

    @property
    def total_time(self) -> float:
        return self.end_time - self.start_time

    @property
    def files_per_second(self) -> float:
        return self.files_processed / self.total_time if self.total_time > 0 else 0

    @property
    def documents_per_second(self) -> float:
        return self.documents_extracted / self.total_time if self.total_time > 0 else 0

@dataclass
class DocumentContent:
    """
    Represents the content of an extracted document.
    """
    file_path: str
    content_type: str
    content: str
    metadata: Dict[str, Any]
    extracted_at: datetime
    file_size: int
    processing_time: float

@dataclass
class Document:
    """A document for the RAG System."""
    id: str
    content: str
    metadata: Dict[str, Any]
    source: str
    chunk_index: int = 0
    embedding: Optional[List[float]] = None
    created_at: datetime = None
    
    def __post_init__(self):
        """Initializes the created_at timestamp if it's not set."""
        if self.created_at is None:
            self.created_at = datetime.now()

class IntelligentFileProcessor:
    """
    An intelligent file processing system.
    """

    def __init__(self, config: Dict[str, Any] = None):
        """
        Initializes the IntelligentFileProcessor.
        """
        self.config = config or {}
        self.supported_extensions = {
            # Text & Code Files
            '.txt': 'text', '.md': 'text', '.rst': 'text', '.log': 'text',
            '.py': 'code', '.js': 'code', '.ts': 'code', '.jsx': 'code', '.tsx': 'code',
            '.java': 'code', '.cpp': 'code', '.c': 'code', '.h': 'code', '.hpp': 'code',
            '.cs': 'code', '.php': 'code', '.rb': 'code', '.go': 'code', '.rs': 'code',
            '.swift': 'code', '.kt': 'code', '.scala': 'code', '.dart': 'code',
            '.html': 'code', '.css': 'code', '.scss': 'code', '.sass': 'code',
            '.xml': 'code', '.json': 'code', '.yaml': 'code', '.yml': 'code',
            '.toml': 'code', '.ini': 'code', '.cfg': 'code', '.conf': 'code',
            '.sh': 'code', '.bat': 'code', '.ps1': 'code', '.sql': 'code',

            # Document Files
            '.pdf': 'document', '.docx': 'document', '.doc': 'document',
            '.xlsx': 'document', '.xls': 'document', '.pptx': 'document',
            '.ppt': 'document', '.odt': 'document', '.ods': 'document',

            # Data Files
            '.csv': 'data', '.tsv': 'data',

            # Configuration Files
            '.env': 'config', '.gitignore': 'config', '.dockerfile': 'config',
            '.dockerignore': 'config', 'docker-compose.yml': 'config',
            'package.json': 'config', 'requirements.txt': 'config',
            'pom.xml': 'config', 'build.gradle': 'config', 'Cargo.toml': 'config',
            'go.mod': 'config', 'composer.json': 'config', 'Gemfile': 'config'
        }

        self.size_limits = {
            'text': 10 * 1024 * 1024,
            'code': 5 * 1024 * 1024,
            'document': 50 * 1024 * 1024,
            'data': 100 * 1024 * 1024,
            'config': 1 * 1024 * 1024
        }

        self.max_workers = self.config.get('max_workers', min(32, (os.cpu_count() or 1) + 4))
        self.chunk_size = self.config.get('chunk_size', 1000)
        self.chunk_overlap = self.config.get('chunk_overlap', 200)

        logger.info(f"🚀 Intelligent File Processor is ready (Workers: {self.max_workers})")

    async def process_directory_deep(self, root_path: str,
                                   include_patterns: List[str] = None,
                                   exclude_patterns: List[str] = None) -> List[DocumentContent]:
        start_time = time.time()

        try:
            include_patterns = include_patterns or ['*']
            exclude_patterns = exclude_patterns or [
                '*.exe', '*.dll', '*.so', '*.dylib', '*.bin',
                '*.zip', '*.tar', '*.gz', '*.rar', '*.7z',
                '*.mp3', '*.mp4', '*.avi', '*.mov', '*.wmv',
                '*.jpg', '*.jpeg', '*.png', '*.gif', '*.bmp',
                '*.iso', '*.img', '*.vmdk', '*.vhd'
            ]

            all_files = self._find_files_recursive(root_path, include_patterns, exclude_patterns)
            logger.info(f"🔍 Found {len(all_files)} files in {root_path}")

            documents = await self._process_files_parallel(all_files)

            end_time = time.time()
            processing_time = end_time - start_time

            logger.info(f"✅ Processing complete: {len(documents)} documents in {processing_time:.2f} seconds")
            logger.info(f"📊 Speed: {len(documents)/processing_time:.2f} documents/second")

            return documents

        except Exception as e:
            logger.error(f"❌ Error processing directory: {e}")
            return []

    def _find_files_recursive(self, root_path: str, include_patterns: List[str],
                            exclude_patterns: List[str]) -> List[str]:
        files = []
        root_path = Path(root_path)

        try:
            for pattern in include_patterns:
                for file_path in root_path.rglob(pattern):
                    if file_path.is_file():
                        should_exclude = False
                        for exclude_pattern in exclude_patterns:
                            if file_path.match(exclude_pattern):
                                should_exclude = True
                                break

                        if not should_exclude:
                            files.append(str(file_path))

            return list(set(files))

        except Exception as e:
            logger.error(f"❌ Error finding files: {e}")
            return []

    async def _process_files_parallel(self, file_paths: List[str]) -> List[DocumentContent]:
        documents = []
        batch_size = max(1, len(file_paths) // self.max_workers)
        batches = [file_paths[i:i + batch_size] for i in range(0, len(file_paths), batch_size)]

        tasks = [asyncio.create_task(self._process_file_batch(batch)) for batch in batches]
        batch_results = await asyncio.gather(*tasks, return_exceptions=True)

        for batch_result in batch_results:
            if isinstance(batch_result, list):
                documents.extend(batch_result)
            else:
                logger.error(f"❌ Batch processing error: {batch_result}")

        return documents

    async def _process_file_batch(self, file_paths: List[str]) -> List[DocumentContent]:
        documents = []
        for file_path in file_paths:
            try:
                doc = await self._process_single_file(file_path)
                if doc:
                    documents.append(doc)
            except Exception as e:
                logger.warning(f"⚠️ Could not process file {file_path}: {e}")
                continue
        return documents

    async def _process_single_file(self, file_path: str) -> Optional[DocumentContent]:
        start_time = time.time()
        try:
            file_path_obj = Path(file_path)
            if not file_path_obj.exists(): return None

            file_size = file_path_obj.stat().st_size
            if file_size == 0: return None

            content_type = self._determine_content_type(file_path_obj)
            if not content_type: return None

            size_limit = self.size_limits.get(content_type, 1 * 1024 * 1024)
            if file_size > size_limit:
                logger.info(f"⚠️ Skipping file {file_path} (too large: {file_size} bytes)")
                return None

            content = await self._extract_content(file_path_obj, content_type)
            if not content: return None

            metadata = self._create_metadata(file_path_obj, content_type, file_size)
            processing_time = time.time() - start_time

            return DocumentContent(
                file_path=str(file_path_obj),
                content_type=content_type,
                content=content,
                metadata=metadata,
                extracted_at=datetime.now(),
                file_size=file_size,
                processing_time=processing_time
            )
        except Exception as e:
            logger.error(f"❌ Error processing file {file_path}: {e}")
            return None

    def _determine_content_type(self, file_path: Path) -> Optional[str]:
        file_name = file_path.name.lower()
        for ext, content_type in self.supported_extensions.items():
            if file_name.endswith(ext):
                return content_type
        for pattern, content_type in self.supported_extensions.items():
            if pattern in file_name:
                return content_type
        try:
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type:
                if mime_type.startswith('text/'): return 'text'
                elif mime_type.startswith('application/'): return 'document'
        except: pass
        return None

    async def _extract_content(self, file_path: Path, content_type: str) -> Optional[str]:
        try:
            if content_type == 'text': return await self._extract_text_content(file_path)
            elif content_type == 'code': return await self._extract_code_content(file_path)
            elif content_type == 'document': return await self._extract_document_content(file_path)
            elif content_type == 'data': return await self._extract_data_content(file_path)
            elif content_type == 'config': return await self._extract_config_content(file_path)
            else: return None
        except Exception as e:
            logger.error(f"❌ Error extracting content from {file_path}: {e}")
            return None

    async def _extract_text_content(self, file_path: Path) -> Optional[str]:
        try:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError: continue
                with open(file_path, 'rb') as f:
                    return f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"❌ Could not read text file {file_path}: {e}")
            return None

    async def _extract_code_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None

            code_metadata = self._analyze_code_structure(content, file_path.suffix)
            enhanced_content = f"// File: {file_path.name}\n"
            enhanced_content += f"// Language: {code_metadata.get('language', 'unknown')}\n"
            enhanced_content += f"// Functions: {len(code_metadata.get('functions', []))}\n"
            enhanced_content += f"// Classes: {len(code_metadata.get('classes', []))}\n"
            enhanced_content += "// Content:\n" + content
            return enhanced_content
        except Exception as e:
            logger.error(f"❌ Could not extract code content from {file_path}: {e}")
            return None

    def _analyze_code_structure(self, content: str, file_extension: str) -> Dict[str, Any]:
        metadata = {'language': self._get_language_from_extension(file_extension), 'functions': [], 'classes': [], 'imports': [], 'comments': []}
        try:
            if file_extension == '.py' and AST_AVAILABLE:
                try:
                    tree = ast.parse(content)
                    for node in ast.walk(tree):
                        if isinstance(node, ast.FunctionDef): metadata['functions'].append(node.name)
                        elif isinstance(node, ast.ClassDef): metadata['classes'].append(node.name)
                        elif isinstance(node, ast.Import):
                            for alias in node.names: metadata['imports'].append(alias.name)
                        elif isinstance(node, ast.ImportFrom):
                            if node.module: metadata['imports'].append(node.module)
                except: pass

            comment_patterns = {'.py': [r'#.*$', r'""".*?"""', r"'''.*?'''"], '.js': [r'//.*$', r'/\*.*?\*/'], '.ts': [r'//.*$', r'/\*.*?\*/'], '.java': [r'//.*$', r'/\*.*?\*/'], '.cpp': [r'//.*$', r'/\*.*?\*/'], '.c': [r'//.*$', r'/\*.*?\*/'], '.html': [r'<!--.*?-->'], '.css': [r'/\*.*?\*/'], '.sql': [r'--.*$', r'/\*.*?\*/']}
            patterns = comment_patterns.get(file_extension, [])
            for pattern in patterns:
                metadata['comments'].extend(re.findall(pattern, content, re.MULTILINE | re.DOTALL))
        except Exception as e:
            logger.warning(f"⚠️ Could not analyze code structure: {e}")
        return metadata

    def _get_language_from_extension(self, extension: str) -> str:
        lang_map = {'.py': 'Python', '.js': 'JavaScript', '.ts': 'TypeScript', '.jsx': 'React JSX', '.tsx': 'React TSX', '.java': 'Java', '.cpp': 'C++', '.c': 'C', '.h': 'C Header', '.hpp': 'C++ Header', '.cs': 'C#', '.php': 'PHP', '.rb': 'Ruby', '.go': 'Go', '.rs': 'Rust', '.swift': 'Swift', '.kt': 'Kotlin', '.scala': 'Scala', '.dart': 'Dart', '.html': 'HTML', '.css': 'CSS', '.scss': 'SCSS', '.sass': 'Sass', '.xml': 'XML', '.json': 'JSON', '.yaml': 'YAML', '.yml': 'YAML', '.toml': 'TOML', '.ini': 'INI', '.cfg': 'Config', '.conf': 'Config', '.sh': 'Shell', '.bat': 'Batch', '.ps1': 'PowerShell', '.sql': 'SQL'}
        return lang_map.get(extension, 'Unknown')

    def _create_metadata(self, file_path: Path, content_type: str, file_size: int) -> Dict[str, Any]:
        try:
            stat = file_path.stat()
            return {'file_name': file_path.name, 'file_path': str(file_path), 'content_type': content_type, 'file_size': file_size, 'file_extension': file_path.suffix.lower(), 'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(), 'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(), 'accessed_time': datetime.fromtimestamp(stat.st_atime).isoformat(), 'is_hidden': file_path.name.startswith('.'), 'parent_directory': str(file_path.parent), 'depth_level': len(file_path.parts) - 1}
        except Exception as e:
            logger.error(f"❌ Could not create metadata for {file_path}: {e}")
            return {}

    async def _extract_document_content(self, file_path: Path) -> Optional[str]:
        try:
            extension = file_path.suffix.lower()
            if extension == '.pdf' and PDF_AVAILABLE: return await self._extract_pdf_content(file_path)
            elif extension == '.docx' and DOCX_AVAILABLE: return await self._extract_docx_content(file_path)
            elif extension in ['.xlsx', '.xls'] and EXCEL_AVAILABLE: return await self._extract_excel_content(file_path)
            elif extension in ['.html', '.htm'] and HTML_AVAILABLE: return await self._extract_html_content(file_path)
            else: return await self._extract_text_content(file_path)
        except Exception as e:
            logger.error(f"❌ Could not extract document content from {file_path}: {e}")
            return None

    async def _extract_pdf_content(self, file_path: Path) -> Optional[str]:
        try:
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                return await loop.run_in_executor(executor, self._extract_pdf_sync, file_path)
        except Exception as e:
            logger.error(f"❌ Could not extract PDF content from {file_path}: {e}")
            return None

    def _extract_pdf_sync(self, file_path: Path) -> Optional[str]:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
        except Exception as e:
            logger.error(f"❌ PDF extraction error: {e}")
            return None

    async def _extract_docx_content(self, file_path: Path) -> Optional[str]:
        try:
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                return await loop.run_in_executor(executor, self._extract_docx_sync, file_path)
        except Exception as e:
            logger.error(f"❌ Could not extract DOCX content from {file_path}: {e}")
            return None

    def _extract_docx_sync(self, file_path: Path) -> Optional[str]:
        try:
            doc = docx.Document(file_path)
            return "".join(p.text + "\n" for p in doc.paragraphs)
        except Exception as e:
            logger.error(f"❌ DOCX extraction error: {e}")
            return None

    async def _extract_excel_content(self, file_path: Path) -> Optional[str]:
        try:
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                return await loop.run_in_executor(executor, self._extract_excel_sync, file_path)
        except Exception as e:
            logger.error(f"❌ Could not extract Excel content from {file_path}: {e}")
            return None

    def _extract_excel_sync(self, file_path: Path) -> Optional[str]:
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content += f"Sheet: {sheet_name}\n"
                for row in sheet.iter_rows(values_only=True):
                    content += "\t".join(str(cell) if cell is not None else "" for cell in row) + "\n"
                content += "\n"
            return content
        except Exception as e:
            logger.error(f"❌ Excel extraction error: {e}")
            return None

    async def _extract_html_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]): script.decompose()
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return ' '.join(chunk for chunk in chunks if chunk)
        except Exception as e:
            logger.error(f"❌ Could not extract HTML content from {file_path}: {e}")
            return None

    async def _extract_data_content(self, file_path: Path) -> Optional[str]:
        try:
            extension = file_path.suffix.lower()
            if extension == '.csv': return await self._extract_csv_content(file_path)
            elif extension == '.json': return await self._extract_json_content(file_path)
            elif extension == '.xml': return await self._extract_xml_content(file_path)
            else: return await self._extract_text_content(file_path)
        except Exception as e:
            logger.error(f"❌ Could not extract data content from {file_path}: {e}")
            return None

    async def _extract_csv_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None
            lines = content.split('\n')
            if lines:
                headers = lines[0].split(',')
                return f"CSV File: {file_path.name}\nColumns: {len(headers)}\nHeaders: {', '.join(headers)}\nRows: {len(lines) - 1}\nContent:\n" + '\n'.join(lines[:100])
            return content
        except Exception as e:
            logger.error(f"❌ Could not extract CSV content from {file_path}: {e}")
            return None

    async def _extract_json_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None
            try:
                data = json.loads(content)
                new_content = f"JSON File: {file_path.name}\nType: {type(data).__name__}\n"
                if isinstance(data, dict): new_content += f"Keys: {', '.join(data.keys())}\n"
                elif isinstance(data, list): new_content += f"Items: {len(data)}\n"
                return new_content + "Content:\n" + content
            except json.JSONDecodeError: return content
        except Exception as e:
            logger.error(f"❌ Could not extract JSON content from {file_path}: {e}")
            return None

    async def _extract_xml_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None
            soup = BeautifulSoup(content, 'xml')
            return f"XML File: {file_path.name}\nRoot Element: {soup.find().name if soup.find() else 'None'}\nContent:\n" + soup.get_text()
        except Exception as e:
            logger.error(f"❌ Could not extract XML content from {file_path}: {e}")
            return None

    async def _extract_config_content(self, file_path: Path) -> Optional[str]:
        try:
            content = await self._extract_text_content(file_path)
            if not content: return None
            return f"Config File: {file_path.name}\nType: Configuration\nContent:\n" + content
        except Exception as e:
            logger.error(f"❌ Could not extract config content from {file_path}: {e}")
            return None

@dataclass
class SearchResult:
    """A search result from the RAG System."""
    document: Document
    score: float
    relevance: str  # 'high', 'medium', 'low'

@dataclass
class RAGResponse:
    """A response from the RAG System."""
    answer: str
    sources: List[Document]
    confidence: float
    context_used: str
    processing_time: float

class EmbeddingProvider:
    """A provider for Embedding Models that uses the UnifiedAIClient."""

    def __init__(self, provider_type: str = "ollama", model_name: str = "nomic-embed-text"):
        """
        Initializes the EmbeddingProvider.

        Args:
            provider_type (str, optional): The type of embedding provider. Defaults to "ollama".
            model_name (str, optional): The name of the embedding model. Defaults to "nomic-embed-text".
        """
        self.provider_type = provider_type
        self.model_name = model_name
        self.ai_client = get_client()
        if not self.ai_client.get_provider(self.provider_type):
             logger.warning(f"⚠️ Provider '{self.provider_type}' not available in UnifiedAIClient.")
             self.is_ready = False
        else:
            self.is_ready = True
            logger.info(f"✅ EmbeddingProvider initialized for provider '{self.provider_type}' with model '{self.model_name}'.")

    async def test_connection(self) -> bool:
        """
        Tests the connection to the embedding service via the UnifiedAIClient.

        Returns:
            bool: True if the connection is successful, False otherwise.
        """
        if not self.is_ready:
            return False
        try:
            # A simple embedding call to test the connection.
            result = self.ai_client.embed(self.provider_type, "test", model=self.model_name)
            return result.get('success', False)
        except Exception as e:
            logger.error(f"❌ Connection test failed for {self.provider_type}: {e}")
            return False

    async def get_embedding(self, text: str) -> Optional[List[float]]:
        """
        Creates an embedding for a text using the UnifiedAIClient.

        Args:
            text (str): The text to create an embedding for.

        Returns:
            Optional[List[float]]: The created embedding, or None if an error occurred.
        """
        if not self.is_ready:
            logger.error(f"❌ Embedding provider '{self.provider_type}' is not ready.")
            return None

        try:
            result = self.ai_client.embed(self.provider_type, text, model=self.model_name)
            if result and result.get('success'):
                return result.get('embedding')
            else:
                logger.error(f"❌ Error creating embedding with {self.provider_type}: {result.get('error')}")
                return None
        except Exception as e:
            logger.error(f"❌ Exception during embedding: {e}")
            return None

class VectorDatabase:
    """Vector Database Manager."""
    
    def __init__(self, db_type: str = "chromadb", config: Dict[str, Any] = None):
        """
        Initializes the VectorDatabase.

        Args:
            db_type (str, optional): The type of vector database. Defaults to "chromadb".
            config (Dict[str, Any], optional): The configuration for the vector database. Defaults to None.
        """
        self.db_type = db_type
        self.config = config or {}
        self.client = None
        self.collection = None
        self.setup_database()
    
    def setup_database(self):
        """Sets up the vector database."""
        try:
            if self.db_type == "chromadb":
                self._setup_chromadb()
            elif self.db_type == "pinecone":
                self._setup_pinecone()
            elif self.db_type == "sqlite":
                self._setup_sqlite_vector()
            else:
                logger.warning(f"⚠️ Vector database not supported: {self.db_type}")
                
        except Exception as e:
            logger.error(f"❌ Error setting up vector database: {e}")
    
    def _setup_chromadb(self):
        """Sets up ChromaDB."""
        if not CHROMADB_AVAILABLE:
            logger.warning("⚠️ ChromaDB not available")
            return
            
        try:
            self.client = chromadb.PersistentClient(
                path=self.config.get("path", "./chroma_db"),
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            
            collection_name = self.config.get("collection_name", "synapse_documents")
            self.collection = self.client.get_or_create_collection(
                name=collection_name,
                metadata={"description": "Synapse RAG Documents"}
            )
            
            logger.info("✅ ChromaDB is ready")
            
        except Exception as e:
            logger.error(f"❌ ChromaDB setup error: {e}")
    
    def _setup_pinecone(self):
        """Sets up Pinecone."""
        if not PINECONE_AVAILABLE:
            logger.warning("⚠️ Pinecone not available")
            return
            
        try:
            api_key = self.config.get("api_key")
            environment = self.config.get("environment")
            
            if not api_key or not environment:
                logger.warning("⚠️ Pinecone API key and environment are required")
                return
            
            pinecone.init(api_key=api_key, environment=environment)
            index_name = self.config.get("index_name", "synapse-rag")
            
            # Create index if it doesn't exist
            if index_name not in pinecone.list_indexes():
                pinecone.create_index(
                    name=index_name,
                    dimension=self.config.get("dimension", 1536),
                    metric="cosine"
                )
            
            self.collection = pinecone.Index(index_name)
            logger.info("✅ Pinecone is ready")
            
        except Exception as e:
            logger.error(f"❌ Pinecone setup error: {e}")
    
    def _setup_sqlite_vector(self):
        """Sets up SQLite for vector storage."""
        try:
            db_path = self.config.get("db_path", "./synapse_vectors.db")
            self.client = sqlite3.connect(db_path)
            
            # Create table for storing vectors
            self.client.execute("""
                CREATE TABLE IF NOT EXISTS document_vectors (
                    id TEXT PRIMARY KEY,
                    content TEXT NOT NULL,
                    embedding TEXT NOT NULL,
                    metadata TEXT,
                    source TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create index for searching
            self.client.execute("""
                CREATE INDEX IF NOT EXISTS idx_document_source 
                ON document_vectors(source)
            """)
            
            self.client.commit()
            logger.info("✅ SQLite Vector Database is ready")
            
        except Exception as e:
            logger.error(f"❌ SQLite vector setup error: {e}")
    
    async def test_connection(self) -> bool:
        """
        Tests the connection to the vector database.

        Returns:
            bool: True if the connection is successful, False otherwise.
        """
        try:
            if self.db_type == "chromadb":
                return self.client is not None and self.collection is not None
            elif self.db_type == "pinecone":
                return self.collection is not None
            elif self.db_type == "sqlite":
                return self.client is not None
            else:
                return False
        except Exception as e:
            logger.error(f"❌ Vector database connection test failed: {e}")
            return False
    
    async def add_documents(self, documents: List[Document]) -> bool:
        """
        Adds documents to the vector database.

        Args:
            documents (List[Document]): A list of documents to add.

        Returns:
            bool: True if the documents were added successfully, False otherwise.
        """
        if not self.collection and not self.client:
            logger.error("❌ Vector database not available")
            return False
        
        try:
            if self.db_type == "chromadb":
                return await self._add_to_chromadb(documents)
            elif self.db_type == "pinecone":
                return await self._add_to_pinecone(documents)
            elif self.db_type == "sqlite":
                return await self._add_to_sqlite(documents)
            else:
                return False
                
        except Exception as e:
            logger.error(f"❌ Error adding documents: {e}")
            return False
    
    async def _add_to_chromadb(self, documents: List[Document]) -> bool:
        """
        Adds documents to ChromaDB.

        Args:
            documents (List[Document]): A list of documents to add.

        Returns:
            bool: True if the documents were added successfully, False otherwise.
        """
        try:
            ids = [doc.id for doc in documents]
            contents = [doc.content for doc in documents]
            embeddings = [doc.embedding for doc in documents if doc.embedding]
            metadatas = [asdict(doc.metadata) for doc in documents]
            
            if embeddings:
                self.collection.add(
                    ids=ids,
                    documents=contents,
                    embeddings=embeddings,
                    metadatas=metadatas
                )
            else:
                self.collection.add(
                    ids=ids,
                    documents=contents,
                    metadatas=metadatas
                )
            
            logger.info(f"✅ Added {len(documents)} documents to ChromaDB")
            return True
            
        except Exception as e:
            logger.error(f"❌ ChromaDB add error: {e}")
            return False
    
    async def _add_to_pinecone(self, documents: List[Document]) -> bool:
        """
        Adds documents to Pinecone.

        Args:
            documents (List[Document]): A list of documents to add.

        Returns:
            bool: True if the documents were added successfully, False otherwise.
        """
        try:
            vectors = []
            for doc in documents:
                if doc.embedding:
                    vectors.append({
                        "id": doc.id,
                        "values": doc.embedding,
                        "metadata": {
                            "content": doc.content,
                            "source": doc.source,
                            **doc.metadata
                        }
                    })
            
            if vectors:
                self.collection.upsert(vectors=vectors)
                logger.info(f"✅ Added {len(vectors)} documents to Pinecone")
                return True
            else:
                return False
                
        except Exception as e:
            logger.error(f"❌ Pinecone add error: {e}")
            return False
    
    async def _add_to_sqlite(self, documents: List[Document]) -> bool:
        """
        Adds documents to SQLite.

        Args:
            documents (List[Document]): A list of documents to add.

        Returns:
            bool: True if the documents were added successfully, False otherwise.
        """
        try:
            cursor = self.client.cursor()
            
            for doc in documents:
                embedding_json = json.dumps(doc.embedding) if doc.embedding else "[]"
                metadata_json = json.dumps(doc.metadata)
                
                cursor.execute("""
                    INSERT OR REPLACE INTO document_vectors 
                    (id, content, embedding, metadata, source)
                    VALUES (?, ?, ?, ?, ?)
                """, (doc.id, doc.content, embedding_json, metadata_json, doc.source))
            
            self.client.commit()
            logger.info(f"✅ Added {len(documents)} documents to SQLite")
            return True
            
        except Exception as e:
            logger.error(f"❌ SQLite add error: {e}")
            return False
    
    async def search(self, query_embedding: List[float], top_k: int = 5) -> List[SearchResult]:
        """
        Searches for similar documents.

        Args:
            query_embedding (List[float]): The embedding of the query.
            top_k (int, optional): The number of results to return. Defaults to 5.

        Returns:
            List[SearchResult]: A list of search results.
        """
        if not self.collection and not self.client:
            logger.error("❌ Vector database not available")
            return []
        
        try:
            if self.db_type == "chromadb":
                return await self._search_chromadb(query_embedding, top_k)
            elif self.db_type == "pinecone":
                return await self._search_pinecone(query_embedding, top_k)
            elif self.db_type == "sqlite":
                return await self._search_sqlite(query_embedding, top_k)
            else:
                return []
                
        except Exception as e:
            logger.error(f"❌ Error during search: {e}")
            return []
    
    async def _search_chromadb(self, query_embedding: List[float], top_k: int) -> List[SearchResult]:
        """
        Searches in ChromaDB.

        Args:
            query_embedding (List[float]): The embedding of the query.
            top_k (int): The number of results to return.

        Returns:
            List[SearchResult]: A list of search results.
        """
        try:
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k
            )
            
            search_results = []
            for i in range(len(results['ids'][0])):
                doc_id = results['ids'][0][i]
                content = results['documents'][0][i]
                metadata = results['metadatas'][0][i]
                distance = results['distances'][0][i]
                
                # Convert distance to score (0-1)
                score = 1 - distance
                
                document = Document(
                    id=doc_id,
                    content=content,
                    metadata=metadata,
                    source=metadata.get('source', 'unknown')
                )
                
                search_results.append(SearchResult(
                    document=document,
                    score=score,
                    relevance=self._get_relevance(score)
                ))
            
            return search_results
            
        except Exception as e:
            logger.error(f"❌ ChromaDB search error: {e}")
            return []
    
    async def _search_pinecone(self, query_embedding: List[float], top_k: int) -> List[SearchResult]:
        """
        Searches in Pinecone.

        Args:
            query_embedding (List[float]): The embedding of the query.
            top_k (int): The number of results to return.

        Returns:
            List[SearchResult]: A list of search results.
        """
        try:
            results = self.collection.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True
            )
            
            search_results = []
            for match in results['matches']:
                metadata = match['metadata']
                
                document = Document(
                    id=match['id'],
                    content=metadata.get('content', ''),
                    metadata=metadata,
                    source=metadata.get('source', 'unknown')
                )
                
                search_results.append(SearchResult(
                    document=document,
                    score=match['score'],
                    relevance=self._get_relevance(match['score'])
                ))
            
            return search_results
            
        except Exception as e:
            logger.error(f"❌ Pinecone search error: {e}")
            return []
    
    async def _search_sqlite(self, query_embedding: List[float], top_k: int) -> List[SearchResult]:
        """
        Searches in SQLite (cosine similarity).

        Args:
            query_embedding (List[float]): The embedding of the query.
            top_k (int): The number of results to return.

        Returns:
            List[SearchResult]: A list of search results.
        """
        try:
            cursor = self.client.cursor()
            cursor.execute("SELECT id, content, embedding, metadata, source FROM document_vectors")
            
            results = []
            query_embedding_np = np.array(query_embedding)
            
            for row in cursor.fetchall():
                doc_id, content, embedding_json, metadata_json, source = row
                
                try:
                    embedding = json.loads(embedding_json)
                    if embedding:
                        embedding_np = np.array(embedding)
                        
                        # Calculate cosine similarity
                        similarity = np.dot(query_embedding_np, embedding_np) / (
                            np.linalg.norm(query_embedding_np) * np.linalg.norm(embedding_np)
                        )
                        
                        metadata = json.loads(metadata_json) if metadata_json else {}
                        
                        document = Document(
                            id=doc_id,
                            content=content,
                            metadata=metadata,
                            source=source
                        )
                        
                        results.append(SearchResult(
                            document=document,
                            score=float(similarity),
                            relevance=self._get_relevance(float(similarity))
                        ))
                        
                except Exception as e:
                    logger.warning(f"⚠️ Could not process document {doc_id}: {e}")
                    continue
            
            # Sort results by score and select top_k
            results.sort(key=lambda x: x.score, reverse=True)
            return results[:top_k]
            
        except Exception as e:
            logger.error(f"❌ SQLite search error: {e}")
            return []
    
    def _get_relevance(self, score: float) -> str:
        """
        Determines the relevance level based on the score.

        Args:
            score (float): The similarity score.

        Returns:
            str: The relevance level ('high', 'medium', or 'low').
        """
        if score >= 0.8:
            return "high"
        elif score >= 0.6:
            return "medium"
        else:
            return "low"

class DocumentProcessor:
    """Document Processing & Chunking"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initializes the DocumentProcessor.

        Args:
            chunk_size (int, optional): The size of each chunk. Defaults to 1000.
            chunk_overlap (int, optional): The overlap between chunks. Defaults to 200.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_document(self, content: str, metadata: Dict[str, Any], source: str) -> List[Document]:
        """
        Processes a document and splits it into chunks.

        Args:
            content (str): The content of the document.
            metadata (Dict[str, Any]): The metadata of the document.
            source (str): The source of the document.

        Returns:
            List[Document]: A list of document chunks.
        """
        try:
            # Clean content
            cleaned_content = self._clean_content(content)
            
            # Split into chunks
            chunks = self._create_chunks(cleaned_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_id = self._generate_doc_id(source, i, chunk)
                
                document = Document(
                    id=doc_id,
                    content=chunk,
                    metadata={
                        **metadata,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "chunk_size": len(chunk)
                    },
                    source=source,
                    chunk_index=i
                )
                
                documents.append(document)
            
            logger.info(f"✅ Processed document {source} into {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error processing document: {e}")
            return []
    
    def _clean_content(self, content: str) -> str:
        """
        Cleans the content.

        Args:
            content (str): The content to clean.

        Returns:
            str: The cleaned content.
        """
        # Remove unnecessary whitespace
        content = " ".join(content.split())
        
        # Remove unwanted special characters
        import re
        content = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]', '', content)
        
        return content.strip()
    
    def _create_chunks(self, content: str) -> List[str]:
        """
        Splits content into chunks.

        Args:
            content (str): The content to split.

        Returns:
            List[str]: A list of content chunks.
        """
        if len(content) <= self.chunk_size:
            return [content]
        
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + self.chunk_size
            
            # Find a suitable split point (sentence or paragraph)
            if end < len(content):
                # Find the nearest sentence end
                sentence_end = content.rfind('.', start, end)
                paragraph_end = content.rfind('\n\n', start, end)
                
                if paragraph_end > sentence_end and paragraph_end > start:
                    end = paragraph_end
                elif sentence_end > start:
                    end = sentence_end + 1
            
            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move to the next chunk (with overlap)
            start = end - self.chunk_overlap
            if start >= len(content):
                break
        
        return chunks
    
    def _generate_doc_id(self, source: str, chunk_index: int, content: str) -> str:
        """
        Generates an ID for a document.

        Args:
            source (str): The source of the document.
            chunk_index (int): The index of the chunk.
            content (str): The content of the chunk.

        Returns:
            str: The generated document ID.
        """
        # Use a hash of the source, chunk index, and content
        content_hash = hashlib.md5(f"{source}_{chunk_index}_{content}".encode()).hexdigest()
        return f"doc_{content_hash[:12]}"

class RAGSystem:
    """
    The main Retrieval-Augmented Generation (RAG) System.
    This class orchestrates the entire RAG pipeline, from document processing
    and embedding to searching and response generation.
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        """
        Initializes the RAGSystem.
        """
        self.config = config or {}
        
        # Initialize components
        self.embedding_provider = EmbeddingProvider(
            provider_type=self.config.get("embedding_provider", "ollama"),
            model_name=self.config.get("embedding_model", "nomic-embed-text")
        )
        self.vector_db = VectorDatabase(
            db_type=self.config.get("vector_db", "chromadb"),
            config=self.config.get("vector_db_config", {})
        )
        self.document_processor = DocumentProcessor(
            chunk_size=self.config.get("chunk_size", 1000),
            chunk_overlap=self.config.get("chunk_overlap", 200)
        )
        self.file_processor = IntelligentFileProcessor(
            config=self.config.get('file_processor', {})
        )

        # Performance monitoring
        self.metrics = PerformanceMetrics(
            start_time=0, end_time=0, files_processed=0, documents_extracted=0,
            embeddings_generated=0, cache_hits=0, cache_misses=0, errors=[]
        )
        
        # Cache system
        self.cache = redis.Redis(
            host=self.config.get("redis_host", "localhost"),
            port=self.config.get("redis_port", 6379),
            db=self.config.get("redis_db", 0),
            decode_responses=True
        )
        self.ai_client = get_client()
        
        logger.info("🚀 Unified RAG System is ready")

    async def scan_directory_deep(self, root_path: str,
                                include_patterns: List[str] = None,
                                exclude_patterns: List[str] = None) -> Dict[str, Any]:
        """
        Scans a directory deeply, processes all files, and adds them to the RAG system.
        """
        start_time = time.time()
        self.metrics.start_time = start_time
        logger.info(f"🔍 Starting deep scan of directory: {root_path}")

        try:
            processed_files = await self.file_processor.process_directory_deep(
                root_path, include_patterns, exclude_patterns
            )
            self.metrics.files_processed = len(processed_files)

            for doc_content in processed_files:
                success = await self.add_document(
                    content=doc_content.content,
                    metadata=doc_content.metadata,
                    source=doc_content.file_path
                )
                if success:
                    self.metrics.embeddings_generated += 1

            end_time = time.time()
            self.metrics.end_time = end_time
            self.metrics.documents_extracted = len(processed_files)

            report = self._create_scan_report(processed_files)
            logger.info(f"✅ Scan complete: {len(processed_files)} documents processed in {end_time - start_time:.2f} seconds")
            return report

        except Exception as e:
            logger.error(f"❌ Error during deep scan: {e}")
            self.metrics.errors.append(str(e))
            return {'error': str(e)}

    def _create_scan_report(self, documents: List[DocumentContent]) -> Dict[str, Any]:
        try:
            content_types = {}
            file_extensions = {}
            total_size = 0

            for doc in documents:
                content_type = doc.content_type
                content_types[content_type] = content_types.get(content_type, 0) + 1

                ext = doc.metadata.get('file_extension', 'unknown')
                file_extensions[ext] = file_extensions.get(ext, 0) + 1

                total_size += doc.file_size

            performance = {
                'total_time': self.metrics.total_time,
                'files_per_second': self.metrics.files_per_second,
                'documents_per_second': self.metrics.documents_per_second,
                'embeddings_generated': self.metrics.embeddings_generated,
                'cache_hits': self.metrics.cache_hits,
                'cache_misses': self.metrics.cache_misses
            }

            return {
                'summary': {
                    'total_documents': len(documents),
                    'total_size_mb': total_size / (1024 * 1024),
                    'content_types': content_types,
                    'file_extensions': dict(sorted(file_extensions.items(), key=lambda x: x[1], reverse=True)[:20]),
                    'performance': performance
                },
                'documents': [
                    {
                        'file_path': doc.file_path,
                        'content_type': doc.content_type,
                        'file_size': doc.file_size
                    }
                    for doc in documents[:100]
                ],
                'errors': self.metrics.errors
            }
        except Exception as e:
            logger.error(f"❌ Error creating report: {e}")
            return {'error': str(e)}

    def get_performance_metrics(self) -> Dict[str, Any]:
        return {
            'total_time': self.metrics.total_time,
            'files_processed': self.metrics.files_processed,
            'cache_hit_rate': self.metrics.cache_hits / (self.metrics.cache_hits + self.metrics.cache_misses) if (self.metrics.cache_hits + self.metrics.cache_misses) > 0 else 0,
            'errors': self.metrics.errors
        }
    
    async def initialize_system(self) -> Dict[str, Any]:
        """
        Initializes the system and tests connections.

        Returns:
            Dict[str, Any]: The status of the system.
        """
        status = {
            "embedding_provider": False,
            "vector_database": False,
            "cache_system": False,
            "system_ready": False
        }
        
        try:
            # Test embedding provider
            logger.info("🔍 Testing Embedding Provider connection...")
            status["embedding_provider"] = await self.embedding_provider.test_connection()
            
            # Test vector database
            logger.info("🔍 Testing Vector Database connection...")
            status["vector_database"] = await self.vector_db.test_connection()
            
            # Test cache system
            logger.info("🔍 Testing Cache System connection...")
            try:
                self.cache.ping()
                status["cache_system"] = True
                logger.info("✅ Cache System is ready")
            except Exception as e:
                logger.warning(f"⚠️ Cache System not available: {e}")
            
            # Check if the system is ready
            status["system_ready"] = status["embedding_provider"] and status["vector_database"]
            
            if status["system_ready"]:
                logger.info("🎉 RAG System is fully operational!")
            else:
                logger.warning("⚠️ RAG System is partially operational")
            
            return status
            
        except Exception as e:
            logger.error(f"❌ Error initializing system: {e}")
            return status
    
    async def get_system_status(self) -> Dict[str, Any]:
        """
        Gets the system status.

        Returns:
            Dict[str, Any]: The status of the system.
        """
        try:
            status = {
                "embedding_provider": {
                    "type": self.embedding_provider.provider_type,
                    "model": self.embedding_provider.model_name,
                    "connected": await self.embedding_provider.test_connection()
                },
                "vector_database": {
                    "type": self.vector_db.db_type,
                    "connected": await self.vector_db.test_connection()
                },
                "cache_system": {
                    "connected": False
                },
                "document_processor": {
                    "chunk_size": self.document_processor.chunk_size,
                    "chunk_overlap": self.document_processor.chunk_overlap
                }
            }
            
            # Test cache
            try:
                self.cache.ping()
                status["cache_system"]["connected"] = True
            except:
                pass
            
            return status
            
        except Exception as e:
            logger.error(f"❌ Could not get system status: {e}")
            return {}
    
    async def add_document(self, content: str, metadata: Dict[str, Any], source: str) -> bool:
        """
        Adds a document to the RAG system.

        Args:
            content (str): The content of the document.
            metadata (Dict[str, Any]): The metadata of the document.
            source (str): The source of the document.

        Returns:
            bool: True if the document was added successfully, False otherwise.
        """
        try:
            # Process the document
            documents = self.document_processor.process_document(content, metadata, source)
            
            if not documents:
                return False
            
            # Create embeddings
            for doc in documents:
                doc.embedding = await self.embedding_provider.get_embedding(doc.content)
            
            # Add to vector database
            success = await self.vector_db.add_documents(documents)
            
            if success:
                # Cache metadata
                cache_key = f"doc_meta:{source}"
                self.cache.setex(cache_key, 3600, json.dumps(metadata))
                
                logger.info(f"✅ Added document {source} successfully ({len(documents)} chunks)")
            
            return success
            
        except Exception as e:
            logger.error(f"❌ Error adding document: {e}")
            return False
    
    async def search(self, query: str, top_k: int = 5, use_cache: bool = True) -> List[SearchResult]:
        """
        Searches for relevant documents.

        Args:
            query (str): The search query.
            top_k (int, optional): The number of results to return. Defaults to 5.
            use_cache (bool, optional): Whether to use the cache. Defaults to True.

        Returns:
            List[SearchResult]: A list of search results.

        Example:
            >>> rag = RAGSystem()
            >>> results = await rag.search("What is the Synapse architecture?")
            >>> for result in results:
            ...     print(f"Source: {result.document.source}, Score: {result.score}")
        """
        try:
            # Check cache
            if use_cache:
                cache_key = f"search:{hashlib.md5(query.encode()).hexdigest()}"
                cached_result = self.cache.get(cache_key)
                if cached_result:
                    logger.info("✅ Using result from cache")
                    return [SearchResult(**json.loads(item)) for item in json.loads(cached_result)]
            
            # Create query embedding
            query_embedding = await self.embedding_provider.get_embedding(query)
            if not query_embedding:
                logger.error("❌ Could not create query embedding")
                return []
            
            # Search in vector database
            results = await self.vector_db.search(query_embedding, top_k)
            
            # Cache results
            if use_cache and results:
                cache_key = f"search:{hashlib.md5(query.encode()).hexdigest()}"
                cache_data = json.dumps([asdict(result) for result in results])
                self.cache.setex(cache_key, 1800, cache_data)  # Cache for 30 minutes
            
            return results
            
        except Exception as e:
            logger.error(f"❌ Error during search: {e}")
            return []
    
    async def generate_response(self, query: str, context_documents: List[Document], 
                              llm_provider: str = "ollama") -> RAGResponse:
        """
        Generates a response using RAG.

        Args:
            query (str): The user's query.
            context_documents (List[Document]): A list of context documents.
            llm_provider (str, optional): The LLM provider to use. Defaults to "ollama".

        Returns:
            RAGResponse: The generated response.

        Example:
            >>> search_results = await rag.search("What is Synapse?")
            >>> documents = [res.document for res in search_results]
            >>> response = await rag.generate_response("What is the Synapse architecture?", documents)
            >>> print(response.answer)
        """
        try:
            start_time = datetime.now()
            
            # Build context from relevant documents
            context = self._build_context(context_documents)
            
            # Create prompt for LLM
            prompt = self._create_rag_prompt(query, context)
            
            # Call LLM
            answer = await self._call_llm(prompt, llm_provider)
            
            # Calculate confidence score
            confidence = self._calculate_confidence(context_documents)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return RAGResponse(
                answer=answer,
                sources=context_documents,
                confidence=confidence,
                context_used=context[:500] + "..." if len(context) > 500 else context,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"❌ Error generating response: {e}")
            return RAGResponse(
                answer="Sorry, an error occurred while processing your request.",
                sources=[],
                confidence=0.0,
                context_used="",
                processing_time=0.0
            )
    
    def _build_context(self, documents: List[Document]) -> str:
        """
        Builds the context from documents.

        Args:
            documents (List[Document]): A list of context documents.

        Returns:
            str: The built context.
        """
        context_parts = []
        
        for i, doc in enumerate(documents, 1):
            context_parts.append(f"Document {i} (from {doc.source}):\n{doc.content}\n")
        
        return "\n".join(context_parts)
    
    def _create_rag_prompt(self, query: str, context: str) -> str:
        """
        Creates a prompt for RAG.

        Args:
            query (str): The user's query.
            context (str): The context from relevant documents.

        Returns:
            str: The created prompt.
        """
        return f"""
You are an AI Assistant that helps answer questions using the provided documents.

Relevant Information:
{context}

Question: {query}

Please answer the question using only the information from the provided documents. If the information is not available, say that you cannot answer.

Answer:
"""
    
    async def _call_llm(self, prompt: str, provider: str) -> str:
        """
        Calls the LLM using the UnifiedAIClient.

        Args:
            prompt (str): The prompt for the LLM.
            provider (str): The LLM provider to use.

        Returns:
            str: The response from the LLM.
        """
        try:
            if not self.ai_client.get_provider(provider):
                return f"LLM provider '{provider}' is not supported or configured."

            messages = [{"role": "user", "content": prompt}]
            # Use a default model from config if available, otherwise let the strategy decide
            model = self.config.get(f"{provider}_model")
            
            result = self.ai_client.generate_response(provider, messages, model=model)

            if result and result.get('success'):
                return result.get('content', 'No content received.')
            else:
                error_msg = result.get('error', 'An unknown error occurred')
                logger.error(f"❌ LLM call error with {provider}: {error_msg}")
                return f"An error occurred while calling the {provider} LLM."

        except Exception as e:
            logger.error(f"❌ Exception during LLM call: {e}")
            return "An unexpected error occurred while calling the LLM."
    
    def _calculate_confidence(self, documents: List[Document]) -> float:
        """
        Calculates the confidence score.

        Args:
            documents (List[Document]): A list of context documents.

        Returns:
            float: The confidence score.
        """
        if not documents:
            return 0.0
        
        # Calculate based on the number of relevant documents and the diversity of sources
        unique_sources = len(set(doc.source for doc in documents))
        total_docs = len(documents)
        
        # Confidence = (number of documents * diversity of sources) / 10
        confidence = (total_docs * unique_sources) / 10.0
        
        return min(confidence, 1.0)  # Limit to 1.0
    
    async def get_statistics(self) -> Dict[str, Any]:
        """
        Gets statistics of the RAG system.

        Returns:
            Dict[str, Any]: A dictionary of statistics.
        """
        try:
            stats = {
                "total_documents": 0,
                "total_chunks": 0,
                "embedding_provider": self.embedding_provider.provider_type,
                "vector_db_type": self.vector_db.db_type,
                "cache_status": "connected" if self.cache.ping() else "disconnected"
            }
            
            # Get statistics from the vector database
            if self.vector_db.db_type == "chromadb" and self.vector_db.collection:
                stats["total_chunks"] = self.vector_db.collection.count()
            
            return stats
            
        except Exception as e:
            logger.error(f"❌ Error getting statistics: {e}")
            return {"error": str(e)}

# Factory function for creating a RAG System
def create_rag_system(config: Dict[str, Any] = None) -> RAGSystem:
    """
    Creates a RAG System instance.

    Args:
        config (Dict[str, Any], optional): The configuration for the RAG system. Defaults to None.

    Returns:
        RAGSystem: A RAG System instance.
    """
    return RAGSystem(config)

# Example usage
if __name__ == "__main__":
    # Example configuration
    config = {
        "embedding_provider": "ollama",
        "embedding_model": "nomic-embed-text",
        "vector_db": "chromadb",
        "vector_db_config": {
            "path": "./chroma_db",
            "collection_name": "synapse_documents"
        },
        "chunk_size": 1000,
        "chunk_overlap": 200,
        "redis_host": "localhost",
        "redis_port": 6379,
        "ollama_model": "llama3.1:8b"
    }
    
    async def test_rag():
        """Tests the RAG system."""
        rag = create_rag_system(config)
        
        # Add a sample document
        sample_content = """
        The Synapse Backend Monolith is an architecture designed to be the central hub of the entire system.
        This system allows for rapid initial development and allows for important components to be gradually
        separated into Microservices later on.
        
        The system uses Python with FastAPI as its core, as it is a high-performance framework
        that is well-suited for API and AI work.
        """
        
        metadata = {
            "title": "Synapse Architecture",
            "author": "Orion Senior Dev",
            "category": "architecture"
        }
        
        success = await rag.add_document(sample_content, metadata, "architecture_doc")
        print(f"Document added successfully: {success}")
        
        # Test search
        results = await rag.search("Synapse architecture", top_k=3)
        print(f"Found {len(results)} relevant documents:")
        
        for result in results:
            print(f"- {result.document.source}: {result.score:.3f}")
        
        # Test response generation
        if results:
            response = await rag.generate_response(
                "What is the Synapse Backend Monolith?",
                [result.document for result in results]
            )
            print(f"\nAnswer: {response.answer}")
            print(f"Confidence: {response.confidence:.3f}")
            print(f"Processing time: {response.processing_time:.2f} seconds")
    
    # Run test
    asyncio.run(test_rag())